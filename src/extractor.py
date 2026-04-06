
"""
eContract Extractor — parses .docx and .txt files and returns a clean,
structured ContractDocument ready for prompt engineering.

FIXES applied vs original:
  1. _read_txt: try utf-8-sig first so BOM is auto-stripped.
  2. _clean_text: explicit lstrip("\ufeff") as belt-and-suspenders.
  3. _ROLE_RE: named group renamed to "pname" (no ambiguity with concatenation).
  4. _extract_parties: preamble-first "Full Name (Alias)" detection added;
     fallback role regex guard against short/fragment names.
  5. _extract_title: operator precedence bug fixed (added inner parentheses).
  6. Governing law regex: handles verbose multi-clause Delaware-style phrasing.
  7. Expiry date: context-aware lookup near expiry keywords instead of last date.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════════
#  Data model
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class ContractParty:
    role: str
    name: str
    address: Optional[str] = None
    wallet_hint: Optional[str] = None


@dataclass
class ContractClause:
    index: int
    heading: str
    raw_text: str
    clause_type: str = "general"
    amount_eth: Optional[str] = None
    deadline_days: Optional[int] = None
    condition: Optional[str] = None


@dataclass
class ContractDocument:
    title: str
    parties: List[ContractParty]
    clauses: List[ContractClause]
    governing_law: Optional[str] = None
    effective_date: Optional[str] = None
    expiry_date: Optional[str] = None
    currency: str = "ETH"
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════════════════════
#  Text-cleaning helpers
# ═══════════════════════════════════════════════════════════════════════════

_NOISE_PATTERNS = [
    r"Page\s+\d+\s+of\s+\d+",
    r"CONFIDENTIAL\s*[-–—]?\s*DO\s+NOT\s+DISTRIBUTE",
    r"DRAFT\s+ONLY",
    r"^\s*[-_=]{4,}\s*$",
    r"\[SIGNATURE\s+BLOCK\]",
    r"\[INTENTIONALLY\s+LEFT\s+BLANK\]",
    r"www\.[^\s]+",
    r"<[^>]+>",
]
_NOISE_RE = re.compile("|".join(_NOISE_PATTERNS), re.IGNORECASE | re.MULTILINE)


def _normalize_whitespace(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    for src, dst in {
        "\u2013": "-", "\u2014": "-",
        "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"',
        "\u00a0": " ", "\u2022": "*", "\u2026": "...",
    }.items():
        text = text.replace(src, dst)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(ln.rstrip() for ln in text.splitlines()).strip()


def _clean_text(raw: str) -> str:
    raw = raw.lstrip("\ufeff")          # FIX 2: strip any residual BOM
    return _normalize_whitespace(_NOISE_RE.sub("", raw))


# ═══════════════════════════════════════════════════════════════════════════
#  Field extractors
# ═══════════════════════════════════════════════════════════════════════════

_ETH_RE = re.compile(r"0x[0-9a-fA-F]{40}")

_AMOUNT_RE = re.compile(
    r"(?:USD|ETH|USDT|DAI|\$|€|£)?\s*[\d,]+(?:\.\d+)?\s*(?:ETH|USDT|DAI|USD|dollars?|ether)?",
    re.IGNORECASE,
)

_DAYS_RE = re.compile(r"(\d+)\s+(?:calendar\s+)?days?", re.IGNORECASE)

_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|\d{4}[/-]\d{1,2}[/-]\d{1,2}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b",
    re.IGNORECASE,
)

# ── Party regexes ───────────────────────────────────────────────────────────

_ROLE_KEYWORDS = [
    "Buyer", "Seller", "Service Provider", "Client", "Vendor", "Contractor",
    "Employer", "Employee", "Licensor", "Licensee", "Lessor", "Lessee",
    "Borrower", "Lender", "Party A", "Party B", "Owner", "Developer",
]

# FIX 3: group renamed to "pname" — avoids any ambiguity from string concat
_ROLE_RE = re.compile(
    r"(?P<role>" + "|".join(re.escape(kw) for kw in _ROLE_KEYWORDS) + r")"
    r"""[:\s"']*(?P<pname>[A-Z][A-Za-z\s,\.]{2,60}?)"""
    r"""(?=\s*(?:\("|,|;|\n|and\b|or\b|hereinafter|$))""",
    re.IGNORECASE,
)

# FIX 4a: explicit "Full Legal Name ("Alias")" pattern — covers most M&A contracts
_PARTY_DEF_RE = re.compile(
    r'([A-Z][A-Za-z\s]{3,60}?)\s+\("([A-Z][A-Za-z\s]{1,30})"\)',
)

# ── Clause regexes ─────────────────────────────────────────────────────────

_CLAUSE_HEAD_RE = re.compile(
    r"^(?:(?:Article|Section|Clause)\s+)?(?:\d+[\.\d]*\.?|[IVXLC]+\.)\s+(.+)$"
    r"|^([A-Z][A-Z\s]{4,})$",
    re.MULTILINE,
)

_CLAUSE_TYPES = {
    "payment":      re.compile(r"payment|price|fee|amount|invoice|compensation", re.I),
    "penalty":      re.compile(r"penalty|liquidated|damages|breach|default|forfeit", re.I),
    "expiry":       re.compile(r"term|duration|expir|terminat|renew|effectiv", re.I),
    "obligation":   re.compile(r"shall|must|obligat|warrant|guarant|represent|covenant", re.I),
    "dispute":      re.compile(r"dispute|arbitrat|mediat|jurisdiction|govern", re.I),
    "confidential": re.compile(r"confidential|non.?disclos|proprietary|secret", re.I),
    "ip":           re.compile(r"intellectual property|copyright|trademark|patent|licen", re.I),
}


def _classify_clause(text: str) -> str:
    for ctype, pat in _CLAUSE_TYPES.items():
        if pat.search(text):
            return ctype
    return "general"


def _extract_amount(text: str) -> Optional[str]:
    m = _AMOUNT_RE.search(text)
    if m:
        raw = m.group(0).strip()
        if raw and any(c.isdigit() for c in raw):
            return raw
    return None


def _extract_days(text: str) -> Optional[int]:
    m = _DAYS_RE.search(text)
    return int(m.group(1)) if m else None


def _extract_eth_addresses(text: str) -> List[str]:
    return _ETH_RE.findall(text)


def _extract_dates(text: str) -> List[str]:
    return _DATE_RE.findall(text)


# ═══════════════════════════════════════════════════════════════════════════
#  Clause splitting
# ═══════════════════════════════════════════════════════════════════════════

def _split_into_clauses(text: str) -> List[ContractClause]:
    clauses: List[ContractClause] = []
    current_heading = "Preamble"
    current_lines: List[str] = []
    idx = 0

    for line in text.splitlines():
        m = _CLAUSE_HEAD_RE.match(line.strip())
        if m:
            body = "\n".join(current_lines).strip()
            if body:
                clauses.append(ContractClause(
                    index=idx,
                    heading=current_heading,
                    raw_text=body,
                    clause_type=_classify_clause(current_heading + " " + body),
                    amount_eth=_extract_amount(body),
                    deadline_days=_extract_days(body),
                ))
                idx += 1
            current_heading = (m.group(1) or m.group(2) or line).strip()
            current_lines = []
        elif line.strip():
            current_lines.append(line)

    body = "\n".join(current_lines).strip()
    if body:
        clauses.append(ContractClause(
            index=idx,
            heading=current_heading,
            raw_text=body,
            clause_type=_classify_clause(current_heading + " " + body),
            amount_eth=_extract_amount(body),
            deadline_days=_extract_days(body),
        ))

    return clauses if clauses else [
        ContractClause(index=0, heading="Full Contract", raw_text=text, clause_type="general")
    ]


# ═══════════════════════════════════════════════════════════════════════════
#  Party extraction
# ═══════════════════════════════════════════════════════════════════════════

def _extract_parties(text: str) -> List[ContractParty]:
    """
    Three-pass extraction:
      Pass 1 — preamble 'Full Name ("Alias")' definitions  [most accurate]
      Pass 2 — role-keyword regex fallback
      Pass 3 — two anonymous placeholders
    """
    eth_addresses = _extract_eth_addresses(text)

    # Pass 1 — scan first 8 000 chars (covers preamble even after long TOC)
    named: List[ContractParty] = []
    for m in _PARTY_DEF_RE.finditer(text[:8000]):
        full_name = m.group(1).strip()
        alias     = m.group(2).strip()
        if len(full_name) > 5:
            named.append(ContractParty(role=alias, name=full_name))
    if named:
        for i, p in enumerate(named):
            p.wallet_hint = eth_addresses[i] if i < len(eth_addresses) else None
        return named

    # Pass 2 — role-keyword regex
    parties: List[ContractParty] = []
    seen: set = set()
    for m in _ROLE_RE.finditer(text):
        role  = m.group("role").strip().title()
        # FIX 4b: use "pname" — the actual group name in _ROLE_RE
        pname = re.sub(r"\s+", " ", m.group("pname")).strip()
        pname = re.sub(r"[,;\.]+$", "", pname)
        if not pname or len(pname) < 3 or role.lower() in seen:
            continue
        seen.add(role.lower())
        wallet = eth_addresses[len(parties)] if len(parties) < len(eth_addresses) else None
        parties.append(ContractParty(role=role, name=pname, wallet_hint=wallet))
    if parties:
        return parties

    # Pass 3 — fallback
    return [
        ContractParty(role="Party A", name="[Party A Name]"),
        ContractParty(role="Party B", name="[Party B Name]"),
    ]


# ═══════════════════════════════════════════════════════════════════════════
#  Title extraction
# ═══════════════════════════════════════════════════════════════════════════

def _extract_title(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        # FIX 5: original had operator-precedence bug — `and` binds tighter than
        #         the multi-condition `or`, making the isupper() branch always True.
        #         Added inner parentheses to group the two alternatives correctly.
        if stripped and 5 < len(stripped) < 150 and (
            stripped.isupper()
            or re.search(r"agreement|contract|deed|memorandum|mou|nda|sla", stripped, re.I)
        ):
            return stripped.title()
    return "Electronic Contract"


# ═══════════════════════════════════════════════════════════════════════════
#  File readers
# ═══════════════════════════════════════════════════════════════════════════

def _read_docx(path: Path) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_txt(path: Path) -> str:
    # FIX 1: utf-8-sig first — Python auto-strips BOM with this codec
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc).lstrip("\ufeff")
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"Could not decode {path} with any known encoding.")


# ═══════════════════════════════════════════════════════════════════════════
#  Public API
# ═══════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".docx", ".txt"}


def extract_contract(file_path: str | Path) -> ContractDocument:
    """
    Parse an eContract file (.docx or .txt) and return a structured
    ContractDocument.
    """
    path = Path(file_path).resolve()
    if not path.exists():
        raise ValueError(f"File not found: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Accepted: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    raw   = _read_docx(path) if ext == ".docx" else _read_txt(path)
    clean = _clean_text(raw)

    title   = _extract_title(clean)
    parties = _extract_parties(clean)
    clauses = _split_into_clauses(clean)
    dates   = _extract_dates(clean)

    # FIX 6: governing law — handles verbose Delaware-style wording
    gov_match = re.search(
        r"governed\s+by(?:\s+and\s+construed\s+in[^.]{0,80}?in\s+accordance\s+with)?"
        r"(?:\s+the)?\s+Laws?\s+of(?:\s+the\s+State\s+of)?\s+([A-Za-z\s]+?)"
        r"(?:\s+without|\.|,|\n)",
        clean, re.I | re.DOTALL,
    )
    if not gov_match:
        gov_match = re.search(
            r"govern(?:ed|ing)\s+by(?:\s+the\s+laws?\s+of)?\s+([A-Za-z\s]+?)(?:\.|,|\n)",
            clean, re.I,
        )
    governing_law = gov_match.group(1).strip() if gov_match else None

    # FIX 7: expiry date — prefer date near expiry/termination keywords;
    # never return a date earlier than the effective date.
    expiry_date: Optional[str] = None
    effective   = dates[0] if dates else None
    expiry_ctx  = re.search(
        r"(?:expir(?:ation|y)|terminat(?:ion)?\s+date|end\s+of\s+term)"
        r"[^.]{0,200}?(" + _DATE_RE.pattern + r")",
        clean, re.I | re.DOTALL,
    )
    if expiry_ctx and expiry_ctx.group(1) != effective:
        expiry_date = expiry_ctx.group(1)
    elif len(dates) > 1 and dates[-1] != effective:
        # Only use last date if it is LATER than effective date (crude heuristic)
        # — avoids financial-period dates that pre-date the contract itself.
        from_year = int(re.search(r"\d{4}", effective).group()) if effective else 0
        last_year = int(re.search(r"\d{4}", dates[-1]).group()) if dates[-1] else 0
        if last_year >= from_year:
            expiry_date = dates[-1]

    return ContractDocument(
        title=title,
        parties=parties,
        clauses=clauses,
        governing_law=governing_law,
        effective_date=dates[0] if dates else None,
        expiry_date=expiry_date,
        full_text=clean,
        metadata={
            "source_file":         str(path),
            "extension":           ext,
            "char_count":          len(clean),
            "clause_count":        len(clauses),
            "party_count":         len(parties),
            "eth_addresses_found": _extract_eth_addresses(clean),
        },
    )
"""
eContract Extractor — parses .docx and .txt files and returns a clean,
structured ContractDocument ready for prompt engineering.

FIXES applied vs original:
  1. _read_txt: try utf-8-sig first so BOM is auto-stripped.
  2. _clean_text: explicit lstrip("\ufeff") as belt-and-suspenders.
  3. _ROLE_RE: named group renamed to "pname" (no ambiguity with concatenation).
  4. _extract_parties: preamble-first "Full Name (Alias)" detection added;
     fallback role regex guard against short/fragment names.
  5. _extract_title: operator precedence bug fixed (added inner parentheses).
  6. Governing law regex: handles verbose multi-clause Delaware-style phrasing.
  7. Expiry date: context-aware lookup near expiry keywords instead of last date.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ═══════════════════════════════════════════════════════════════════════════
#  Data model
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class ContractParty:
    role: str
    name: str
    address: Optional[str] = None
    wallet_hint: Optional[str] = None


@dataclass
class ContractClause:
    index: int
    heading: str
    raw_text: str
    clause_type: str = "general"
    amount_eth: Optional[str] = None
    deadline_days: Optional[int] = None
    condition: Optional[str] = None


@dataclass
class ContractDocument:
    title: str
    parties: List[ContractParty]
    clauses: List[ContractClause]
    governing_law: Optional[str] = None
    effective_date: Optional[str] = None
    expiry_date: Optional[str] = None
    currency: str = "ETH"
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


# ═══════════════════════════════════════════════════════════════════════════
#  Text-cleaning helpers
# ═══════════════════════════════════════════════════════════════════════════

_NOISE_PATTERNS = [
    r"Page\s+\d+\s+of\s+\d+",
    r"CONFIDENTIAL\s*[-–—]?\s*DO\s+NOT\s+DISTRIBUTE",
    r"DRAFT\s+ONLY",
    r"^\s*[-_=]{4,}\s*$",
    r"\[SIGNATURE\s+BLOCK\]",
    r"\[INTENTIONALLY\s+LEFT\s+BLANK\]",
    r"www\.[^\s]+",
    r"<[^>]+>",
]
_NOISE_RE = re.compile("|".join(_NOISE_PATTERNS), re.IGNORECASE | re.MULTILINE)


def _normalize_whitespace(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    for src, dst in {
        "\u2013": "-", "\u2014": "-",
        "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"',
        "\u00a0": " ", "\u2022": "*", "\u2026": "...",
    }.items():
        text = text.replace(src, dst)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(ln.rstrip() for ln in text.splitlines()).strip()


def _clean_text(raw: str) -> str:
    raw = raw.lstrip("\ufeff")          # FIX 2: strip any residual BOM
    return _normalize_whitespace(_NOISE_RE.sub("", raw))


# ═══════════════════════════════════════════════════════════════════════════
#  Field extractors
# ═══════════════════════════════════════════════════════════════════════════

_ETH_RE = re.compile(r"0x[0-9a-fA-F]{40}")

_AMOUNT_RE = re.compile(
    r"(?:USD|ETH|USDT|DAI|\$|€|£)?\s*[\d,]+(?:\.\d+)?\s*(?:ETH|USDT|DAI|USD|dollars?|ether)?",
    re.IGNORECASE,
)

_DAYS_RE = re.compile(r"(\d+)\s+(?:calendar\s+)?days?", re.IGNORECASE)

_DATE_RE = re.compile(
    r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|\d{4}[/-]\d{1,2}[/-]\d{1,2}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})\b",
    re.IGNORECASE,
)

# ── Party regexes ───────────────────────────────────────────────────────────

_ROLE_KEYWORDS = [
    "Buyer", "Seller", "Service Provider", "Client", "Vendor", "Contractor",
    "Employer", "Employee", "Licensor", "Licensee", "Lessor", "Lessee",
    "Borrower", "Lender", "Party A", "Party B", "Owner", "Developer",
]

# FIX 3: group renamed to "pname" — avoids any ambiguity from string concat
_ROLE_RE = re.compile(
    r"(?P<role>" + "|".join(re.escape(kw) for kw in _ROLE_KEYWORDS) + r")"
    r"""[:\s"']*(?P<pname>[A-Z][A-Za-z\s,\.]{2,60}?)"""
    r"""(?=\s*(?:\("|,|;|\n|and\b|or\b|hereinafter|$))""",
    re.IGNORECASE,
)

# FIX 4a: explicit "Full Legal Name ("Alias")" pattern — covers most M&A contracts
_PARTY_DEF_RE = re.compile(
    r'([A-Z][A-Za-z\s]{3,60}?)\s+\("([A-Z][A-Za-z\s]{1,30})"\)',
)

# ── Clause regexes ─────────────────────────────────────────────────────────

_CLAUSE_HEAD_RE = re.compile(
    r"^(?:(?:Article|Section|Clause)\s+)?(?:\d+[\.\d]*\.?|[IVXLC]+\.)\s+(.+)$"
    r"|^([A-Z][A-Z\s]{4,})$",
    re.MULTILINE,
)

_CLAUSE_TYPES = {
    "payment":      re.compile(r"payment|price|fee|amount|invoice|compensation", re.I),
    "penalty":      re.compile(r"penalty|liquidated|damages|breach|default|forfeit", re.I),
    "expiry":       re.compile(r"term|duration|expir|terminat|renew|effectiv", re.I),
    "obligation":   re.compile(r"shall|must|obligat|warrant|guarant|represent|covenant", re.I),
    "dispute":      re.compile(r"dispute|arbitrat|mediat|jurisdiction|govern", re.I),
    "confidential": re.compile(r"confidential|non.?disclos|proprietary|secret", re.I),
    "ip":           re.compile(r"intellectual property|copyright|trademark|patent|licen", re.I),
}


def _classify_clause(text: str) -> str:
    for ctype, pat in _CLAUSE_TYPES.items():
        if pat.search(text):
            return ctype
    return "general"


def _extract_amount(text: str) -> Optional[str]:
    m = _AMOUNT_RE.search(text)
    if m:
        raw = m.group(0).strip()
        if raw and any(c.isdigit() for c in raw):
            return raw
    return None


def _extract_days(text: str) -> Optional[int]:
    m = _DAYS_RE.search(text)
    return int(m.group(1)) if m else None


def _extract_eth_addresses(text: str) -> List[str]:
    return _ETH_RE.findall(text)


def _extract_dates(text: str) -> List[str]:
    return _DATE_RE.findall(text)


# ═══════════════════════════════════════════════════════════════════════════
#  Clause splitting
# ═══════════════════════════════════════════════════════════════════════════

def _split_into_clauses(text: str) -> List[ContractClause]:
    clauses: List[ContractClause] = []
    current_heading = "Preamble"
    current_lines: List[str] = []
    idx = 0

    for line in text.splitlines():
        m = _CLAUSE_HEAD_RE.match(line.strip())
        if m:
            body = "\n".join(current_lines).strip()
            if body:
                clauses.append(ContractClause(
                    index=idx,
                    heading=current_heading,
                    raw_text=body,
                    clause_type=_classify_clause(current_heading + " " + body),
                    amount_eth=_extract_amount(body),
                    deadline_days=_extract_days(body),
                ))
                idx += 1
            current_heading = (m.group(1) or m.group(2) or line).strip()
            current_lines = []
        elif line.strip():
            current_lines.append(line)

    body = "\n".join(current_lines).strip()
    if body:
        clauses.append(ContractClause(
            index=idx,
            heading=current_heading,
            raw_text=body,
            clause_type=_classify_clause(current_heading + " " + body),
            amount_eth=_extract_amount(body),
            deadline_days=_extract_days(body),
        ))

    return clauses if clauses else [
        ContractClause(index=0, heading="Full Contract", raw_text=text, clause_type="general")
    ]


# ═══════════════════════════════════════════════════════════════════════════
#  Party extraction
# ═══════════════════════════════════════════════════════════════════════════

def _extract_parties(text: str) -> List[ContractParty]:
    """
    Three-pass extraction:
      Pass 1 — preamble 'Full Name ("Alias")' definitions  [most accurate]
      Pass 2 — role-keyword regex fallback
      Pass 3 — two anonymous placeholders
    """
    eth_addresses = _extract_eth_addresses(text)

    # Pass 1 — scan first 8 000 chars (covers preamble even after long TOC)
    named: List[ContractParty] = []
    for m in _PARTY_DEF_RE.finditer(text[:8000]):
        full_name = m.group(1).strip()
        alias     = m.group(2).strip()
        if len(full_name) > 5:
            named.append(ContractParty(role=alias, name=full_name))
    if named:
        for i, p in enumerate(named):
            p.wallet_hint = eth_addresses[i] if i < len(eth_addresses) else None
        return named

    # Pass 2 — role-keyword regex
    parties: List[ContractParty] = []
    seen: set = set()
    for m in _ROLE_RE.finditer(text):
        role  = m.group("role").strip().title()
        # FIX 4b: use "pname" — the actual group name in _ROLE_RE
        pname = re.sub(r"\s+", " ", m.group("pname")).strip()
        pname = re.sub(r"[,;\.]+$", "", pname)
        if not pname or len(pname) < 3 or role.lower() in seen:
            continue
        seen.add(role.lower())
        wallet = eth_addresses[len(parties)] if len(parties) < len(eth_addresses) else None
        parties.append(ContractParty(role=role, name=pname, wallet_hint=wallet))
    if parties:
        return parties

    # Pass 3 — fallback
    return [
        ContractParty(role="Party A", name="[Party A Name]"),
        ContractParty(role="Party B", name="[Party B Name]"),
    ]


# ═══════════════════════════════════════════════════════════════════════════
#  Title extraction
# ═══════════════════════════════════════════════════════════════════════════

def _extract_title(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        # FIX 5: original had operator-precedence bug — `and` binds tighter than
        #         the multi-condition `or`, making the isupper() branch always True.
        #         Added inner parentheses to group the two alternatives correctly.
        if stripped and 5 < len(stripped) < 150 and (
            stripped.isupper()
            or re.search(r"agreement|contract|deed|memorandum|mou|nda|sla", stripped, re.I)
        ):
            return stripped.title()
    return "Electronic Contract"


# ═══════════════════════════════════════════════════════════════════════════
#  File readers
# ═══════════════════════════════════════════════════════════════════════════

def _read_docx(path: Path) -> str:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _read_txt(path: Path) -> str:
    # FIX 1: utf-8-sig first — Python auto-strips BOM with this codec
    for enc in ("utf-8-sig", "utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc).lstrip("\ufeff")
        except (UnicodeDecodeError, LookupError):
            continue
    raise ValueError(f"Could not decode {path} with any known encoding.")


# ═══════════════════════════════════════════════════════════════════════════
#  Public API
# ═══════════════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".docx", ".txt"}


def extract_contract(file_path: str | Path) -> ContractDocument:
    """
    Parse an eContract file (.docx or .txt) and return a structured
    ContractDocument.
    """
    path = Path(file_path).resolve()
    if not path.exists():
        raise ValueError(f"File not found: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Accepted: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    raw   = _read_docx(path) if ext == ".docx" else _read_txt(path)
    clean = _clean_text(raw)

    title   = _extract_title(clean)
    parties = _extract_parties(clean)
    clauses = _split_into_clauses(clean)
    dates   = _extract_dates(clean)

    # FIX 6: governing law — handles verbose Delaware-style wording
    gov_match = re.search(
        r"governed\s+by(?:\s+and\s+construed\s+in[^.]{0,80}?in\s+accordance\s+with)?"
        r"(?:\s+the)?\s+Laws?\s+of(?:\s+the\s+State\s+of)?\s+([A-Za-z\s]+?)"
        r"(?:\s+without|\.|,|\n)",
        clean, re.I | re.DOTALL,
    )
    if not gov_match:
        gov_match = re.search(
            r"govern(?:ed|ing)\s+by(?:\s+the\s+laws?\s+of)?\s+([A-Za-z\s]+?)(?:\.|,|\n)",
            clean, re.I,
        )
    governing_law = gov_match.group(1).strip() if gov_match else None

    # FIX 7: expiry date — prefer date near expiry/termination keywords;
    # never return a date earlier than the effective date.
    expiry_date: Optional[str] = None
    effective   = dates[0] if dates else None
    expiry_ctx  = re.search(
        r"(?:expir(?:ation|y)|terminat(?:ion)?\s+date|end\s+of\s+term)"
        r"[^.]{0,200}?(" + _DATE_RE.pattern + r")",
        clean, re.I | re.DOTALL,
    )
    if expiry_ctx and expiry_ctx.group(1) != effective:
        expiry_date = expiry_ctx.group(1)
    elif len(dates) > 1 and dates[-1] != effective:
        # Only use last date if it is LATER than effective date (crude heuristic)
        # — avoids financial-period dates that pre-date the contract itself.
        from_year = int(re.search(r"\d{4}", effective).group()) if effective else 0
        last_year = int(re.search(r"\d{4}", dates[-1]).group()) if dates[-1] else 0
        if last_year >= from_year:
            expiry_date = dates[-1]

    return ContractDocument(
        title=title,
        parties=parties,
        clauses=clauses,
        governing_law=governing_law,
        effective_date=dates[0] if dates else None,
        expiry_date=expiry_date,
        full_text=clean,
        metadata={
            "source_file":         str(path),
            "extension":           ext,
            "char_count":          len(clean),
            "clause_count":        len(clauses),
            "party_count":         len(parties),
            "eth_addresses_found": _extract_eth_addresses(clean),
        },
    )